"""Institutional DOCX renderer backed by python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from legal_ai.domain.errors import InvalidFinalizationError, RendererExecutionError


class PythonDocxRenderer:
    """Render a serializable canonical snapshot into a temporary DOCX."""

    name = "python-docx"
    version = "1.1"
    timeout_seconds = 30

    def render(self, snapshot: dict[str, Any], output_path: Path) -> None:
        try:
            from docx import Document
            from docx.enum.section import WD_ORIENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt

            document_data = snapshot.get("document")
            source_text = snapshot.get("source_text")
            if (
                not isinstance(document_data, dict)
                or not isinstance(source_text, str)
                or not source_text.strip()
            ):
                raise InvalidFinalizationError(details={"field": "document"})
            title = document_data.get("title")
            locale = document_data.get("locale")
            if not isinstance(title, str) or not title.strip():
                raise InvalidFinalizationError(details={"field": "document.title"})
            if locale != "es-AR":
                raise InvalidFinalizationError(details={"field": "document.locale"})

            word = Document()
            section = word.sections[0]
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)

            normal = word.styles["Normal"]
            normal.font.name = "Arial"
            normal.font.size = Pt(11)
            normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

            header = document_data.get("institutional_header", "")
            if header:
                section.header.paragraphs[0].text = str(header)
                self._set_font(section.header.paragraphs[0], "Arial", 11)

            title_paragraph = word.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.add_run(title.strip())
            title_run.bold = True
            title_run.font.name = "Arial"
            title_run.font.size = Pt(12)

            self._add_section(word, "VISTO", document_data.get("visto", []))
            self._add_section(
                word, "CONSIDERANDO", document_data.get("considerando", [])
            )
            por_ello = document_data.get("por_ello", "")
            if por_ello:
                self._add_heading(word, "POR ELLO")
                self._add_body(word, str(por_ello))

            articles = document_data.get("articles", [])
            if isinstance(articles, list):
                for index, article in enumerate(articles, start=1):
                    if isinstance(article, dict):
                        number = article.get("number", index)
                        text = article.get("text", article.get("content", ""))
                    else:
                        number, text = index, article
                    self._add_heading(word, f"ARTÍCULO {number}°")
                    self._add_body(word, str(text))

            signatures = document_data.get("signatures", [])
            if isinstance(signatures, list):
                for signature in signatures:
                    paragraph = word.add_paragraph()
                    paragraph.paragraph_format.space_before = Pt(72)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(str(signature))
                    run.font.name = "Arial"
                    run.font.size = Pt(11)

            if document_data.get("page_numbers") is True:
                footer = section.footer.paragraphs[0]
                footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = footer.add_run()
                field_begin = OxmlElement("w:fldChar")
                field_begin.set(qn("w:fldCharType"), "begin")
                instruction = OxmlElement("w:instrText")
                instruction.set(qn("xml:space"), "preserve")
                instruction.text = " PAGE "
                field_end = OxmlElement("w:fldChar")
                field_end.set(qn("w:fldCharType"), "end")
                run._r.extend((field_begin, instruction, field_end))

            properties = word.core_properties
            properties.title = title.strip()
            properties.author = None
            properties.last_modified_by = None
            output_path.parent.mkdir(parents=True, exist_ok=True)
            word.save(str(output_path))
        except InvalidFinalizationError:
            raise
        except Exception as exc:
            raise RendererExecutionError() from exc

    @staticmethod
    def _set_font(paragraph: Any, name: str, size: int) -> None:
        from docx.shared import Pt

        for run in paragraph.runs:
            run.font.name = name
            run.font.size = Pt(size)

    def _add_heading(self, document: Any, text: str) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)

    def _add_body(self, document: Any, text: str) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)

    def _add_section(self, document: Any, heading: str, values: Any) -> None:
        self._add_heading(document, heading)
        if isinstance(values, list):
            for value in values:
                self._add_body(document, str(value))
